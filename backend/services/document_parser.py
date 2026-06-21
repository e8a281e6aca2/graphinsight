"""Document parser adapters for graph ingestion."""
from __future__ import annotations

import json
import mimetypes
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

import httpx

from config import get_settings
from core import get_logger
from services.runtime_config import get_document_parser_runtime_config


logger = get_logger()
settings = get_settings()


@dataclass
class ParsedBlock:
    text: str
    block_type: str = "text"
    heading_path: List[str] = field(default_factory=list)
    page_start: Optional[int] = None
    page_end: Optional[int] = None
    source_location: str = ""


@dataclass
class ParsedDocument:
    text: str
    parser_provider: str
    parser_version: str = ""
    parse_mode: str = ""
    blocks: List[ParsedBlock] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    raw_payload: Any = None
    raw_output_path: str = ""

    def page_range(self) -> tuple[Optional[int], Optional[int]]:
        starts = [block.page_start for block in self.blocks if block.page_start is not None]
        ends = [block.page_end for block in self.blocks if block.page_end is not None]
        return (min(starts) if starts else None, max(ends) if ends else None)


class DocumentParserError(RuntimeError):
    """Raised when a parser cannot produce text for a document."""


class NativeDocumentParser:
    provider = "native"
    version = "builtin-pdfplumber-pypdf-docx"
    parse_mode = "builtin"

    def parse(self, path: Path) -> ParsedDocument:
        text, error = self._read_text(path)
        warnings = [error] if error else []
        blocks = [
            ParsedBlock(
                text=text,
                block_type="text",
                source_location=path.name,
            )
        ] if text.strip() else []
        return ParsedDocument(
            text=text,
            parser_provider=self.provider,
            parser_version=self.version,
            parse_mode=self.parse_mode,
            blocks=blocks,
            warnings=warnings,
        )

    def _read_text(self, path: Path) -> tuple[str, str | None]:
        ext = path.suffix.lower()
        if ext in {".txt", ".md", ".markdown", ".csv", ".log"}:
            return path.read_text(encoding="utf-8", errors="ignore"), None
        if ext == ".json":
            try:
                content = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
                return json.dumps(content, ensure_ascii=False, indent=2), None
            except Exception:
                return path.read_text(encoding="utf-8", errors="ignore"), None
        if ext == ".docx":
            try:
                import docx  # type: ignore
            except Exception:
                logger.warning("缺少 python-docx，无法解析 docx", context={"file": str(path)})
                return "", "missing_python_docx"
            doc = docx.Document(str(path))
            return "\n".join([p.text for p in doc.paragraphs if p.text]), None
        if ext == ".pdf":
            try:
                import pdfplumber  # type: ignore

                with pdfplumber.open(str(path)) as pdf:
                    pages = [(page.extract_text() or "") for page in pdf.pages]
                return "\n".join(pages), None
            except Exception as exc:
                logger.warning(
                    "pdfplumber 解析失败，回退 pypdf",
                    context={"file": str(path), "error": str(exc)},
                )

            try:
                from pypdf import PdfReader  # type: ignore
            except Exception:
                logger.warning("缺少 pypdf，无法解析 pdf", context={"file": str(path)})
                return "", "missing_pypdf"
            try:
                reader = PdfReader(str(path))
                pages = []
                for page in reader.pages:
                    text = page.extract_text() or ""
                    pages.append(text)
                return "\n".join(pages), None
            except Exception as exc:  # noqa: BLE001
                logger.warning("PDF 解析失败，已跳过", context={"file": str(path), "error": str(exc)})
                return "", f"pdf_parse_error: {exc}"
        return "", "unsupported_file"


class MinerUDocumentParser:
    provider = "mineru"

    def __init__(self, config: Dict[str, Any]) -> None:
        self.config = config

    def parse(self, path: Path) -> ParsedDocument:
        base_url = str(self.config.get("base_url") or "").strip().rstrip("/")
        if not base_url:
            raise DocumentParserError("mineru_base_url_missing")
        endpoint_path = str(self.config.get("endpoint_path") or "/file_parse").strip() or "/file_parse"
        if not endpoint_path.startswith("/"):
            endpoint_path = f"/{endpoint_path}"
        url = f"{base_url}{endpoint_path}"
        parse_mode = str(self.config.get("parse_mode") or "auto").strip() or "auto"
        timeout = float(self.config.get("timeout_seconds") or 300)
        output_format = str(self.config.get("output_format") or "markdown,json").strip()
        file_field = str(self.config.get("file_field") or "files").strip() or "files"
        mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"

        data = {
            "parse_method": parse_mode,
            "parse_mode": parse_mode,
            "output_format": output_format,
        }
        with path.open("rb") as file_obj:
            files = {file_field: (path.name, file_obj, mime_type)}
            with httpx.Client(timeout=timeout, trust_env=settings.http_client_trust_env) as client:
                response = client.post(url, data=data, files=files)
        if response.status_code >= 400:
            raise DocumentParserError(f"mineru_http_{response.status_code}: {response.text[:300]}")

        payload = self._decode_response(response)
        text = self._extract_text(payload)
        blocks = self._extract_blocks(payload)
        if not text and blocks:
            text = "\n\n".join(block.text for block in blocks if block.text.strip())
        if not text.strip():
            raise DocumentParserError("mineru_empty_text")

        version = self._first_text(payload, ["parser_version", "version", "mineru_version"]) or str(
            self.config.get("parser_version") or ""
        )
        warnings = self._extract_warnings(payload)
        return ParsedDocument(
            text=text,
            parser_provider=self.provider,
            parser_version=version,
            parse_mode=parse_mode,
            blocks=blocks or [ParsedBlock(text=text, source_location=path.name)],
            warnings=warnings,
            raw_payload=payload,
        )

    @staticmethod
    def _decode_response(response: httpx.Response) -> Any:
        content_type = response.headers.get("content-type", "")
        if "json" in content_type.lower():
            return response.json()
        try:
            return response.json()
        except Exception:
            return response.text

    @classmethod
    def _extract_text(cls, payload: Any) -> str:
        if isinstance(payload, str):
            return payload
        direct = cls._first_text(
            payload,
            [
                "markdown",
                "md",
                "md_content",
                "content",
                "text",
                "plain_text",
                "result_text",
            ],
        )
        if direct:
            return direct
        if isinstance(payload, dict):
            for key in ("results", "data", "result", "output", "document"):
                if key in payload:
                    nested = cls._extract_text(payload[key])
                    if nested:
                        return nested
            metadata_keys = {
                "task_id",
                "status",
                "backend",
                "file_names",
                "created_at",
                "started_at",
                "completed_at",
                "error",
                "status_url",
                "result_url",
                "version",
                "protocol_version",
                "queued_tasks",
                "processing_tasks",
                "completed_tasks",
                "failed_tasks",
            }
            for key, value in payload.items():
                if key in metadata_keys:
                    continue
                if isinstance(value, (dict, list)):
                    nested = cls._extract_text(value)
                    if nested:
                        return nested
        if isinstance(payload, list):
            texts = []
            for item in payload:
                nested = cls._extract_text(item)
                if nested:
                    texts.append(nested)
            return "\n\n".join(texts).strip()
        return ""

    @classmethod
    def _extract_blocks(cls, payload: Any) -> List[ParsedBlock]:
        candidates: List[Any] = []
        if isinstance(payload, dict):
            for key in ("blocks", "content_list", "pages", "paragraphs", "items"):
                value = payload.get(key)
                if isinstance(value, list):
                    candidates.extend(value)
            for key in ("data", "result", "output", "document"):
                if isinstance(payload.get(key), (dict, list)):
                    candidates.extend(cls._extract_blocks(payload[key]))
        elif isinstance(payload, list):
            candidates.extend(payload)

        blocks: List[ParsedBlock] = []
        for item in candidates:
            if isinstance(item, ParsedBlock):
                blocks.append(item)
                continue
            if isinstance(item, str):
                text = item.strip()
                if text:
                    blocks.append(ParsedBlock(text=text))
                continue
            if not isinstance(item, dict):
                continue
            text = cls._first_text(item, ["text", "content", "markdown", "md", "value"])
            if not text:
                continue
            page = cls._first_int(item, ["page", "page_no", "page_idx", "page_start"])
            page_end = cls._first_int(item, ["page_end", "end_page"]) or page
            block_type = cls._first_text(item, ["block_type", "type", "category"]) or "text"
            headings = item.get("heading_path") or item.get("headings") or []
            if isinstance(headings, str):
                heading_path = [part.strip() for part in headings.split("/") if part.strip()]
            elif isinstance(headings, list):
                heading_path = [str(part).strip() for part in headings if str(part).strip()]
            else:
                heading_path = []
            source_location = cls._first_text(item, ["source_location", "location", "bbox"]) or (
                f"Page {page}" if page is not None else ""
            )
            blocks.append(
                ParsedBlock(
                    text=text,
                    block_type=block_type,
                    heading_path=heading_path,
                    page_start=page,
                    page_end=page_end,
                    source_location=source_location,
                )
            )
        return blocks

    @classmethod
    def _extract_warnings(cls, payload: Any) -> List[str]:
        warnings = []
        if isinstance(payload, dict):
            value = payload.get("warnings") or payload.get("warning")
            if isinstance(value, list):
                warnings.extend(str(item) for item in value if str(item).strip())
            elif value:
                warnings.append(str(value))
            for key in ("data", "result", "output"):
                if key in payload:
                    warnings.extend(cls._extract_warnings(payload[key]))
        return warnings[:20]

    @classmethod
    def _first_text(cls, payload: Any, keys: Iterable[str]) -> str:
        if not isinstance(payload, dict):
            return ""
        for key in keys:
            value = payload.get(key)
            if isinstance(value, str) and value.strip():
                return value.strip()
            if isinstance(value, (int, float)):
                return str(value)
        return ""

    @staticmethod
    def _first_int(payload: Dict[str, Any], keys: Iterable[str]) -> Optional[int]:
        for key in keys:
            value = payload.get(key)
            if value is None:
                continue
            try:
                return int(value)
            except Exception:
                continue
        return None


class DocumentParserManager:
    def __init__(self, config: Optional[Dict[str, Any]] = None) -> None:
        self.config = config or get_document_parser_runtime_config()
        self.native = NativeDocumentParser()

    def parse(self, path: Path, provider_override: Optional[str] = None) -> ParsedDocument:
        provider = self._normalize_provider(provider_override or self.config.get("provider") or "native")
        candidates = self._candidate_providers(path, provider)
        failures: List[str] = []
        for candidate in candidates:
            try:
                parsed = self._parser(candidate).parse(path)
                if failures:
                    parsed.warnings = [f"fallback_from_{provider}: {'; '.join(failures)}"] + parsed.warnings
                return parsed
            except Exception as exc:  # noqa: BLE001
                reason = f"{candidate}: {exc}"
                failures.append(reason)
                logger.warning("文档解析器失败", context={"file": str(path), "provider": candidate, "error": str(exc)})
        raise DocumentParserError("; ".join(failures) if failures else "no_parser_available")

    def _candidate_providers(self, path: Path, provider: str) -> List[str]:
        if provider == "native":
            return ["native"]
        fallback = self._normalize_provider(self.config.get("fallback_provider") or "native", allow_none=True)
        if provider == "mineru" and path.suffix.lower() != ".pdf":
            return ["native"]
        candidates = [provider]
        if fallback != "none" and fallback not in candidates:
            candidates.append(fallback)
        return candidates

    def _parser(self, provider: str):
        if provider == "native":
            return self.native
        if provider == "mineru":
            return MinerUDocumentParser(self.config)
        raise DocumentParserError(f"unsupported_parser_provider: {provider}")

    @staticmethod
    def _normalize_provider(value: Any, *, allow_none: bool = False) -> str:
        provider = str(value or "native").strip().lower()
        if allow_none and provider == "none":
            return "none"
        return provider if provider in {"native", "mineru"} else "native"


__all__ = [
    "DocumentParserError",
    "DocumentParserManager",
    "MinerUDocumentParser",
    "NativeDocumentParser",
    "ParsedBlock",
    "ParsedDocument",
]
